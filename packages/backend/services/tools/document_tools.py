"""Document generation and export tools.

generate_document: Create PDF/DOCX from structured content.
export_transcript: Export transcript in TXT/SRT/VTT/DOCX/PDF format.
"""

from __future__ import annotations

import logging
import os
import uuid
from pathlib import Path

from services.tool_registry import Artifact, ToolContext, ToolDef, ToolResult

logger = logging.getLogger(__name__)

# Directory for generated files (served by the API)
_GENERATED_DIR = str(Path(__file__).resolve().parent.parent.parent / "generated_documents")


def _ensure_dir():
    os.makedirs(_GENERATED_DIR, exist_ok=True)


async def handle_generate_document(args: dict, ctx: ToolContext) -> ToolResult:
    """Generate a PDF or DOCX document from structured content."""
    title = args.get("title", "Document")
    fmt = args.get("format", "pdf").lower()
    sections = args.get("sections", [])

    if not sections:
        return ToolResult(content="No sections provided. Please specify at least one section with a heading and content.")

    _ensure_dir()
    safe_title = "".join(c if c.isalnum() or c in " -_" else "" for c in title).strip()[:50]
    filename = f"{safe_title}_{uuid.uuid4().hex[:8]}.{fmt}"
    filepath = os.path.join(_GENERATED_DIR, filename)

    try:
        if fmt == "pdf":
            _generate_pdf(filepath, title, sections)
        elif fmt == "docx":
            _generate_docx(filepath, title, sections)
        else:
            return ToolResult(content=f"Unsupported format: {fmt}. Use 'pdf' or 'docx'.")
    except Exception as e:
        logger.exception("Document generation failed")
        return ToolResult(content=f"Failed to generate document: {e}")

    return ToolResult(
        content=f"Created {fmt.upper()} document: {title}",
        artifacts=[Artifact(
            type="file_download",
            data={"url": f"/api/ai/generated/{filename}", "filename": filename},
        )],
    )


def _generate_pdf(filepath: str, title: str, sections: list[dict]):
    """Generate a PDF using reportlab."""
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer

    doc = SimpleDocTemplate(filepath, pagesize=letter)
    styles = getSampleStyleSheet()
    story = []

    # Title
    story.append(Paragraph(title, styles["Title"]))
    story.append(Spacer(1, 12))

    for section in sections:
        heading = section.get("heading", "")
        content = section.get("content", "")
        if heading:
            story.append(Paragraph(heading, styles["Heading2"]))
            story.append(Spacer(1, 6))
        for para in content.split("\n"):
            if para.strip():
                story.append(Paragraph(para, styles["Normal"]))
                story.append(Spacer(1, 4))
        story.append(Spacer(1, 8))

    doc.build(story)


def _generate_docx(filepath: str, title: str, sections: list[dict]):
    """Generate a DOCX using python-docx."""
    from docx import Document

    doc = Document()
    doc.add_heading(title, level=0)

    for section in sections:
        heading = section.get("heading", "")
        content = section.get("content", "")
        if heading:
            doc.add_heading(heading, level=2)
        for para in content.split("\n"):
            if para.strip():
                doc.add_paragraph(para)

    doc.save(filepath)


async def handle_export_transcript(args: dict, ctx: ToolContext) -> ToolResult:
    """Export a transcript in the specified format."""
    transcript_id = args.get("transcript_id", "")
    fmt = args.get("format", "txt").lower()

    from sqlalchemy import select
    from sqlalchemy.orm import selectinload
    from persistence.models import Transcript, Segment

    # Verify transcript exists (eager-load recording for title)
    result = await ctx.db.execute(
        select(Transcript)
        .where(Transcript.id == transcript_id)
        .options(selectinload(Transcript.recording))
    )
    transcript = result.scalar_one_or_none()
    if not transcript:
        return ToolResult(content=f"Transcript '{transcript_id}' not found.")

    # Get segments
    seg_result = await ctx.db.execute(
        select(Segment)
        .where(Segment.transcript_id == transcript_id)
        .order_by(Segment.start_time)
    )
    segments = seg_result.scalars().all()

    if not segments:
        return ToolResult(content="Transcript has no segments to export.")

    _ensure_dir()
    recording_title = getattr(transcript, 'recording', None)
    rec_name = recording_title.title if recording_title else "transcript"
    safe_name = "".join(c if c.isalnum() or c in " -_" else "" for c in rec_name).strip()[:40]
    filename = f"{safe_name}_{uuid.uuid4().hex[:8]}.{fmt}"
    filepath = os.path.join(_GENERATED_DIR, filename)

    try:
        if fmt == "txt":
            _export_txt(filepath, segments)
        elif fmt == "srt":
            _export_srt(filepath, segments)
        elif fmt == "vtt":
            _export_vtt(filepath, segments)
        else:
            return ToolResult(content=f"Unsupported export format: {fmt}. Use txt, srt, or vtt.")
    except Exception as e:
        logger.exception("Transcript export failed")
        return ToolResult(content=f"Export failed: {e}")

    return ToolResult(
        content=f"Exported transcript as {fmt.upper()}: {rec_name}",
        artifacts=[Artifact(
            type="file_download",
            data={"url": f"/api/ai/generated/{filename}", "filename": filename},
        )],
    )


def _export_txt(filepath: str, segments):
    """Export as plain text."""
    with open(filepath, "w", encoding="utf-8") as f:
        for seg in segments:
            speaker = getattr(seg, "speaker", None) or ""
            if speaker:
                f.write(f"[{speaker}] ")
            f.write(f"{seg.text}\n")


def _export_srt(filepath: str, segments):
    """Export as SRT subtitle format."""
    def _ts(seconds: float) -> str:
        h = int(seconds // 3600)
        m = int((seconds % 3600) // 60)
        s = int(seconds % 60)
        ms = int((seconds % 1) * 1000)
        return f"{h:02d}:{m:02d}:{s:02d},{ms:03d}"

    with open(filepath, "w", encoding="utf-8") as f:
        for i, seg in enumerate(segments, 1):
            start = getattr(seg, "start_time", 0.0) or 0.0
            end = getattr(seg, "end_time", start + 1.0) or start + 1.0
            f.write(f"{i}\n")
            f.write(f"{_ts(start)} --> {_ts(end)}\n")
            f.write(f"{seg.text}\n\n")


def _export_vtt(filepath: str, segments):
    """Export as WebVTT format."""
    def _ts(seconds: float) -> str:
        h = int(seconds // 3600)
        m = int((seconds % 3600) // 60)
        s = int(seconds % 60)
        ms = int((seconds % 1) * 1000)
        return f"{h:02d}:{m:02d}:{s:02d}.{ms:03d}"

    with open(filepath, "w", encoding="utf-8") as f:
        f.write("WEBVTT\n\n")
        for seg in segments:
            start = getattr(seg, "start_time", 0.0) or 0.0
            end = getattr(seg, "end_time", start + 1.0) or start + 1.0
            f.write(f"{_ts(start)} --> {_ts(end)}\n")
            f.write(f"{seg.text}\n\n")


generate_document_tool = ToolDef(
    name="generate_document",
    description="Create a downloadable PDF or DOCX document from structured content.",
    parameters={
        "type": "object",
        "properties": {
            "title": {"type": "string", "description": "Document title"},
            "format": {"type": "string", "enum": ["pdf", "docx"], "description": "Output format"},
            "sections": {
                "type": "array",
                "items": {
                    "type": "object",
                    "properties": {
                        "heading": {"type": "string"},
                        "content": {"type": "string"},
                    },
                },
                "description": "Document sections with headings and content",
            },
        },
        "required": ["title", "format", "sections"],
    },
    handler=handle_generate_document,
    project_scoped=False,
)

export_transcript_tool = ToolDef(
    name="export_transcript",
    description="Export an attached transcript as TXT, SRT, or VTT.",
    parameters={
        "type": "object",
        "properties": {
            "transcript_id": {"type": "string", "description": "The transcript ID to export"},
            "format": {"type": "string", "enum": ["txt", "srt", "vtt"], "description": "Export format"},
        },
        "required": ["transcript_id", "format"],
    },
    handler=handle_export_transcript,
    project_scoped=True,
)
