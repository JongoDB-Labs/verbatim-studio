#!/usr/bin/env python3
"""Build the tour demo workspace zip.

Reads `tour-data/manifest.json` + `tour-data/files/`, generates the
DOCX/PDF documents from the manifest's `extracted_text` (so the zip
is reproducible from the manifest alone), bundles everything into
`tour-demo.zip`, and writes the result to `dist/tour-demo.zip`.

Build artifacts go under `dist/` (gitignored). The tour-data source
directory is committed; the zip is uploaded to verbatim-studio-releases
as a release asset on each tag push.

Usage:
    python scripts/build_tour_demo.py
    python scripts/build_tour_demo.py --output /tmp/tour-demo.zip
"""

from __future__ import annotations

import argparse
import json
import logging
import shutil
import sys
import zipfile
from pathlib import Path

logger = logging.getLogger(__name__)

REPO_ROOT = Path(__file__).resolve().parents[1]
TOUR_DATA_DIR = REPO_ROOT / "tour-data"
DEFAULT_OUTPUT = REPO_ROOT / "dist" / "tour-demo.zip"


def _generate_pdf(text: str, out_path: Path, title: str) -> None:
    """Render plain text into a basic PDF using reportlab.

    Keeps the layout simple — title heading + body paragraphs. We're
    not building a typesetting masterpiece, just something Verbatim
    can ingest and demonstrate OCR / search against.
    """
    from reportlab.lib.pagesizes import LETTER
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import (
        Paragraph, SimpleDocTemplate, Spacer,
    )

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "TitleStyle",
        parent=styles["Heading1"],
        fontSize=18,
        spaceAfter=20,
    )
    body_style = ParagraphStyle(
        "BodyStyle",
        parent=styles["BodyText"],
        fontSize=11,
        leading=15,
        spaceAfter=10,
    )
    heading_style = ParagraphStyle(
        "HeadingStyle",
        parent=styles["Heading2"],
        fontSize=13,
        spaceAfter=8,
    )

    doc = SimpleDocTemplate(
        str(out_path),
        pagesize=LETTER,
        leftMargin=0.75 * inch,
        rightMargin=0.75 * inch,
        topMargin=0.75 * inch,
        bottomMargin=0.75 * inch,
    )

    flowables: list = [Paragraph(title, title_style), Spacer(1, 0.1 * inch)]
    for paragraph in text.split("\n\n"):
        paragraph = paragraph.strip()
        if not paragraph:
            continue
        # Treat ALL-CAPS short lines as section headings
        first_line = paragraph.split("\n", 1)[0]
        if first_line.isupper() and len(first_line) < 60:
            flowables.append(Paragraph(first_line, heading_style))
            rest = paragraph.split("\n", 1)[1] if "\n" in paragraph else ""
            if rest.strip():
                flowables.append(Paragraph(
                    rest.replace("\n", "<br/>"), body_style,
                ))
        else:
            flowables.append(Paragraph(
                paragraph.replace("\n", "<br/>"), body_style,
            ))

    doc.build(flowables)


def _generate_docx(text: str, out_path: Path, title: str) -> None:
    """Render plain text into a .docx file using python-docx."""
    from docx import Document  # type: ignore
    from docx.shared import Pt

    document = Document()
    heading = document.add_heading(title, level=1)
    heading.style.font.size = Pt(20)

    for paragraph in text.split("\n\n"):
        paragraph = paragraph.strip()
        if not paragraph:
            continue
        first_line = paragraph.split("\n", 1)[0]
        rest = paragraph.split("\n", 1)[1] if "\n" in paragraph else ""
        # Treat ALL-CAPS short lines as section headings
        if first_line.isupper() and len(first_line) < 60:
            document.add_heading(first_line, level=2)
            if rest.strip():
                document.add_paragraph(rest)
        else:
            document.add_paragraph(paragraph)

    document.save(str(out_path))


def _generate_text(text: str, out_path: Path) -> None:
    """Write plain-text content to disk."""
    out_path.write_text(text, encoding="utf-8")


def _generate_image_placeholder(text: str, out_path: Path) -> None:
    """Generate a simple whiteboard-style image with the provided text.

    Used when no real Unsplash photo is available. The build script
    prefers a committed real photo if present; this is the fallback
    so the build never breaks because of a missing asset.
    """
    from PIL import Image, ImageDraw, ImageFont

    img = Image.new("RGB", (1200, 800), color=(245, 245, 245))
    draw = ImageDraw.Draw(img)

    try:
        # Try a system font; falls back to default if unavailable.
        font_path = "/System/Library/Fonts/Helvetica.ttc"
        title_font = ImageFont.truetype(font_path, 36)
        body_font = ImageFont.truetype(font_path, 22)
    except Exception:
        title_font = ImageFont.load_default()
        body_font = ImageFont.load_default()

    y = 60
    for i, line in enumerate(text.split("\n")):
        line = line.rstrip()
        if i == 0:
            draw.text((60, y), line, fill=(30, 30, 30), font=title_font)
            y += 60
        else:
            draw.text((60, y), line, fill=(50, 50, 50), font=body_font)
            y += 36

    img.save(str(out_path), "JPEG", quality=85)


def build(output: Path) -> None:
    """Build the tour demo zip."""
    manifest_path = TOUR_DATA_DIR / "manifest.json"
    if not manifest_path.exists():
        raise SystemExit(f"manifest missing: {manifest_path}")

    manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
    files_dir = TOUR_DATA_DIR / "files"
    files_dir.mkdir(parents=True, exist_ok=True)
    (files_dir / "documents").mkdir(exist_ok=True)
    (files_dir / "images").mkdir(exist_ok=True)
    (files_dir / "audio").mkdir(exist_ok=True)

    # Generate document files from extracted_text in the manifest. We
    # do this at build time rather than committing pre-built binaries
    # for two reasons:
    #   1. text + a generator is cheaper to review than binary .docx
    #   2. text-only diffs make manifest changes visible in code review
    #
    # Real binary assets (audio, photos) MUST be committed under
    # files/ — they can't be reproduced from the manifest.
    for doc in manifest.get("documents", []):
        rel_path = doc["filename"]
        abs_path = files_dir / rel_path
        abs_path.parent.mkdir(parents=True, exist_ok=True)
        text = doc.get("extracted_text", "")
        title = doc.get("title", "")
        mime = doc.get("mime_type", "")

        if mime == "application/pdf":
            logger.info("Generating PDF: %s", rel_path)
            _generate_pdf(text, abs_path, title)
        elif mime == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            logger.info("Generating DOCX: %s", rel_path)
            _generate_docx(text, abs_path, title)
        elif mime == "text/plain":
            logger.info("Writing TXT: %s", rel_path)
            _generate_text(text, abs_path)
        elif mime == "image/jpeg":
            # Image: prefer a committed real photo; fall back to a
            # generated placeholder if none is present.
            if not abs_path.exists() or abs_path.stat().st_size == 0:
                logger.warning(
                    "No real image at %s — generating placeholder. "
                    "Add a real photo for production builds.", rel_path,
                )
                _generate_image_placeholder(text, abs_path)
            else:
                logger.info("Using committed real image: %s", rel_path)
        else:
            logger.warning("Unknown mime type for %s: %s — skipping", rel_path, mime)

    # Audio files MUST exist already (they aren't reproducible). Verify.
    for rec in manifest.get("recordings", []):
        rel_path = rec["filename"]
        abs_path = files_dir / rel_path
        if not abs_path.exists():
            raise SystemExit(
                f"Missing audio file: {abs_path}. "
                f"Audio files must be committed under tour-data/files/."
            )
        if abs_path.stat().st_size == 0:
            raise SystemExit(
                f"Empty audio file: {abs_path}. Re-fetch from source."
            )

    # Build the zip
    output.parent.mkdir(parents=True, exist_ok=True)
    if output.exists():
        output.unlink()
    logger.info("Building zip: %s", output)
    with zipfile.ZipFile(output, "w", zipfile.ZIP_DEFLATED, compresslevel=6) as z:
        z.writestr("manifest.json", json.dumps(manifest, indent=2))
        for path in files_dir.rglob("*"):
            if path.is_file():
                arcname = "files/" + str(path.relative_to(files_dir))
                z.write(path, arcname)

    size_mb = output.stat().st_size / 1024 / 1024
    logger.info("Tour demo zip built: %s (%.1f MB)", output, size_mb)


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--output", type=Path, default=DEFAULT_OUTPUT)
    parser.add_argument("--verbose", "-v", action="store_true")
    args = parser.parse_args(argv)

    logging.basicConfig(
        level=logging.DEBUG if args.verbose else logging.INFO,
        format="%(asctime)s [%(name)s] %(levelname)s: %(message)s",
    )

    build(args.output)
    return 0


if __name__ == "__main__":
    sys.exit(main())
